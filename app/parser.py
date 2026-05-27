import io
import os
import shutil
import pytesseract
import zipfile
import xml.etree.ElementTree as ET
import re
import html
from PIL import Image
from pypdf import PdfReader
from docx import Document

# --- CONFIGURATION ---
# Auto-detect Tesseract OCR binary path across common locations
def find_tesseract_path() -> str:
    path_in_env = shutil.which("tesseract")
    if path_in_env:
        return path_in_env

    standard_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    
    # Check user home folder (e.g. C:\Users\Aryan Panigrahi)
    user_home = os.path.expanduser("~")
    if user_home:
        standard_paths.append(os.path.join(user_home, "tesseract.exe"))
        standard_paths.append(os.path.join(user_home, "Tesseract-OCR", "tesseract.exe"))
        
    local_app_data = os.environ.get("LOCALAPPDATA")
    if local_app_data:
        standard_paths.append(os.path.join(local_app_data, "Tesseract-OCR", "tesseract.exe"))
        
    for path in standard_paths:
        if os.path.exists(path):
            return path
            
    return r"C:\Program Files\Tesseract-OCR\tesseract.exe"

pytesseract.pytesseract.tesseract_cmd = find_tesseract_path()

def parse_pdf_content(file_content: bytes) -> str:
    """
    High-fidelity PDF text extractor using PyMuPDF (fitz) with fallback to pypdf and Tesseract OCR.
    """
    text = ""
    try:
        import fitz
        print("[INFO] PDF detected! Opening with PyMuPDF (fitz) for high-performance extraction...")
        doc = fitz.open(stream=file_content, filetype="pdf")
        page_index = 1
        for page in doc:
            page_text = page.get_text() or ""
            
            # If page contains standard native text, use it directly
            if len(page_text.strip()) >= 50:
                text += page_text + "\n"
            else:
                # If page text is very sparse, render the entire page to a high-res image and run OCR
                print(f"[WARN] Page {page_index} has sparse text. Rendering page to image for Tesseract OCR...")
                try:
                    # Render page at 150 DPI (optimal balance of speed and character recognition accuracy)
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    img_obj = Image.open(io.BytesIO(img_data))
                    
                    ocr_result = pytesseract.image_to_string(img_obj)
                    if ocr_result.strip():
                        text += ocr_result + "\n"
                    else:
                        text += page_text + "\n"
                except Exception as page_ocr_err:
                    print(f"[WARN] Page {page_index} OCR rendering failed: {page_ocr_err}")
                    text += page_text + "\n"
            page_index += 1
    except Exception as fitz_err:
        print(f"[WARN] PyMuPDF failed: {fitz_err}. Falling back to pypdf...")
        # Fallback to pypdf
        reader = PdfReader(io.BytesIO(file_content))
        page_index = 1
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if len(page_text.strip()) >= 50:
                text += page_text + "\n"
            else:
                print(f"[WARN] Page {page_index} has sparse text. Attempting image OCR fallback...")
                ocr_text = ""
                try:
                    if hasattr(page, "images") and page.images:
                        for img in page.images:
                            try:
                                img_obj = Image.open(io.BytesIO(img.data))
                                ocr_result = pytesseract.image_to_string(img_obj)
                                if ocr_result.strip():
                                    ocr_text += ocr_result + "\n"
                            except Exception as img_err:
                                print(f"[WARN] Error parsing image on Page {page_index}: {img_err}")
                    if ocr_text.strip():
                        text += ocr_text
                    else:
                        text += page_text + "\n"
                except Exception as ocr_err:
                    print(f"[WARN] OCR extraction failed on Page {page_index}: {ocr_err}")
                    text += page_text + "\n"
            page_index += 1
    return text.strip()

def extract_printable_strings(data: bytes, min_len: int = 4) -> str:
    """
    Robust binary strings extractor. Scans binary files for sequences of printable characters
    in UTF-16 (little/big endian) and standard ASCII/UTF-8. Highly effective for legacy formats like .doc, .xls, .ppt.
    """
    text_parts = []
    
    # 1. Try decoding UTF-16 strings (extremely common in legacy MS Word and Windows files)
    for encoding in ["utf-16", "utf-16-le", "utf-16-be"]:
        try:
            decoded = data.decode(encoding, errors="ignore")
            blocks = []
            current_block = []
            for char in decoded:
                o = ord(char)
                # Printable ASCII, common punctuation/whitespaces, and standard European extensions
                if (32 <= o <= 126) or char in "\n\r\t" or (160 <= o <= 8225):
                    current_block.append(char)
                else:
                    if len(current_block) >= min_len:
                        blocks.append("".join(current_block))
                    current_block = []
            if len(current_block) >= min_len:
                blocks.append("".join(current_block))
            
            cleaned_utf16 = "\n".join([b.strip() for b in blocks if len(b.strip()) >= min_len])
            if cleaned_utf16.strip():
                text_parts.append(cleaned_utf16.strip())
        except Exception:
            pass

    # 2. Try ASCII extraction
    blocks = []
    current_block = []
    for byte in data:
        if (32 <= byte <= 126) or byte in (10, 13, 9):
            current_block.append(chr(byte))
        else:
            if len(current_block) >= min_len:
                blocks.append("".join(current_block))
            current_block = []
    if len(current_block) >= min_len:
        blocks.append("".join(current_block))
        
    cleaned_ascii = "\n".join([b.strip() for b in blocks if len(b.strip()) >= min_len])
    if cleaned_ascii.strip():
        text_parts.append(cleaned_ascii.strip())
        
    # Choose the extraction method that recovered the largest amount of readable text
    if not text_parts:
        return ""
    return max(text_parts, key=len)

def parse_odt(file_content: bytes) -> str:
    """
    Parses OpenDocument Text (.odt) by unzipping in-memory and reading paragraphs and headings from content.xml.
    """
    with zipfile.ZipFile(io.BytesIO(file_content)) as z:
        content_xml = z.read("content.xml")
        root = ET.fromstring(content_xml)
        text_parts = []
        for elem in root.iter():
            # Check for paragraphs and headers (opendocument text elements namespaces end in text:p or text:h)
            if elem.tag.endswith("}p") or elem.tag.endswith("}h"):
                text_val = "".join(elem.itertext()).strip()
                if text_val:
                    text_parts.append(text_val)
        return "\n".join(text_parts)

def parse_pages(file_content: bytes) -> str:
    """
    Parses Apple Pages (.pages) packages. Checks for QuickLook preview (PDF or JPG) and falls back
    to string extraction on internal database files if a preview is unavailable.
    """
    with zipfile.ZipFile(io.BytesIO(file_content)) as z:
        namelist = z.namelist()
        
        # Modern Pages documents include a preview PDF inside
        preview_pdf_files = [name for name in namelist if "preview.pdf" in name.lower() or "quicklook/preview.pdf" in name.lower()]
        if preview_pdf_files:
            print(f"[INFO] Apple Pages PDF preview found: {preview_pdf_files[0]}. Processing PDF extraction...")
            pdf_bytes = z.read(preview_pdf_files[0])
            return parse_pdf_content(pdf_bytes)
            
        # Check for image previews
        preview_img_files = [name for name in namelist if "preview" in name.lower() and (name.endswith(".jpg") or name.endswith(".jpeg") or name.endswith(".png"))]
        if preview_img_files:
            print(f"[INFO] Apple Pages Image preview found: {preview_img_files[0]}. Processing image OCR...")
            img_bytes = z.read(preview_img_files[0])
            image = Image.open(io.BytesIO(img_bytes))
            return pytesseract.image_to_string(image)
            
        # Fallback: scan binary database archives (.iwa files) inside Pages zip for readable strings
        print("[WARN] No QuickLook previews found inside Pages. Extracting binary strings...")
        fallback_text = []
        for name in namelist:
            if name.endswith(".iwa") or "document" in name.lower():
                try:
                    data = z.read(name)
                    extracted = extract_printable_strings(data)
                    if extracted.strip():
                        fallback_text.append(extracted.strip())
                except Exception:
                    pass
        if fallback_text:
            return "\n".join(fallback_text)
            
        raise ValueError("Could not find any readable text or preview inside the Apple Pages document.")

def parse_rtf(file_content: bytes) -> str:
    """
    Extracts text from Rich Text Format (.rtf) files by stripping out control words and markup.
    """
    try:
        text = file_content.decode('utf-8', errors='ignore')
    except Exception:
        text = file_content.decode('latin-1', errors='ignore')
        
    pattern = re.compile(r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\\'([0-9a-f]{2})|\\([^a-z])|([{}])|([^\\{}]+)", re.IGNORECASE)
    words = []
    for match in pattern.finditer(text):
        word, number, hex_char, special, brace, plain = match.groups()
        if plain:
            words.append(plain)
        elif hex_char:
            try:
                words.append(bytes.fromhex(hex_char).decode('ansi', errors='ignore'))
            except Exception:
                pass
        elif special == '~':
            words.append(' ')
        elif special in ['-', '_']:
            words.append(special)
            
    cleaned = "".join(words)
    cleaned = re.sub(r'[ \t\r\f]+', ' ', cleaned)
    cleaned = re.sub(r'(\s*\n\s*)+', '\n', cleaned)
    
    # Fallback to smart binary string extractor if the resulting parsed text is extremely short
    if len(cleaned.strip()) < 50:
        return extract_printable_strings(file_content)
        
    return cleaned.strip()

def parse_html(file_content: bytes) -> str:
    """
    Strips CSS style blocks, JavaScript, and tags from HTML/XML to return clean plain text.
    """
    try:
        text = file_content.decode('utf-8', errors='ignore')
    except Exception:
        text = file_content.decode('latin-1', errors='ignore')
        
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = html.unescape(text)
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return "\n".join(lines)

def parse_xlsx(file_content: bytes) -> str:
    """
    Extracts text cells and shared strings from Excel Spreadsheet (.xlsx) archives.
    """
    text_parts = []
    with zipfile.ZipFile(io.BytesIO(file_content)) as z:
        namelist = z.namelist()
        if "xl/sharedStrings.xml" in namelist:
            try:
                root = ET.fromstring(z.read("xl/sharedStrings.xml"))
                for elem in root.iter():
                    if elem.tag.endswith("}t"):
                        val = elem.text
                        if val and val.strip():
                            text_parts.append(val.strip())
            except Exception:
                pass
                
        for name in namelist:
            if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"):
                try:
                    root = ET.fromstring(z.read(name))
                    for elem in root.iter():
                        if elem.tag.endswith("}v") or elem.tag.endswith("}t"):
                            val = elem.text
                            if val and val.strip():
                                text_parts.append(val.strip())
                except Exception:
                    pass
    return "\n".join(text_parts)

def parse_pptx(file_content: bytes) -> str:
    """
    Extracts slide text boxes and shapes from PowerPoint Presentation (.pptx) archives.
    """
    text_parts = []
    with zipfile.ZipFile(io.BytesIO(file_content)) as z:
        namelist = z.namelist()
        slide_files = sorted([name for name in namelist if name.startswith("ppt/slides/slide") and name.endswith(".xml")])
        for slide in slide_files:
            try:
                root = ET.fromstring(z.read(slide))
                for elem in root.iter():
                    if elem.tag.endswith("}t"):
                        val = elem.text
                        if val and val.strip():
                            text_parts.append(val.strip())
            except Exception:
                pass
    return "\n".join(text_parts)

async def parse_file(file_content: bytes, filename: str) -> dict:
    """
    Determines file type and extracts text accordingly.
    Supports ALL formats: PDF, DOCX, DOC, Pages, ODT, RTF, XLSX, PPTX, HTML, XML,
    Images (PNG, JPG, WebP, etc. via OCR), Text, Code files, and unknown binary structures.
    """
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    text = ""
    file_type = "text"

    print(f"[INFO] Parsing file: {filename} (extension: {ext})")

    try:
        # 1. PDF
        if ext == 'pdf':
            text = parse_pdf_content(file_content)
            file_type = "pdf"

        # 2. DOCX / DOC
        elif ext in ['docx', 'doc']:
            try:
                doc = Document(io.BytesIO(file_content))
                text_parts = []
                
                # Extract text from sections, headers
                for section in doc.sections:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                text_parts.append(para.text.strip())
                                
                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                        
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_txt = cell.text.strip()
                            if cell_txt and cell_txt not in row_text:
                                row_text.append(cell_txt)
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                            
                text = "\n".join(text_parts)
                file_type = "docx"
            except Exception as docx_err:
                print(f"[WARN] python-docx parsing failed: {docx_err}. Falling back to binary string extraction...")
                # Handles legacy .doc files or corrupted docx archives
                text = extract_printable_strings(file_content)
                file_type = "doc_extracted"

        # 3. OpenDocument Text (ODT)
        elif ext == 'odt':
            try:
                text = parse_odt(file_content)
                file_type = "odt"
            except Exception as odt_err:
                print(f"[WARN] ODT parser failed: {odt_err}. Falling back to binary string extraction...")
                text = extract_printable_strings(file_content)
                file_type = "odt_extracted"

        # 4. Apple Pages (.pages)
        elif ext == 'pages':
            try:
                text = parse_pages(file_content)
                file_type = "pages"
            except Exception as pages_err:
                print(f"[WARN] Pages parser failed: {pages_err}. Falling back to binary string extraction...")
                text = extract_printable_strings(file_content)
                file_type = "pages_extracted"

        # 5. Images (OCR with Binary Fallback)
        elif ext in ['jpg', 'jpeg', 'png', 'webp', 'tiff', 'bmp', 'gif', 'heic', 'heif']:
            print("[INFO] Image file detected. Running Tesseract OCR...")
            try:
                image = Image.open(io.BytesIO(file_content))
                text = pytesseract.image_to_string(image)
                file_type = "image_ocr"
            except Exception as img_err:
                print(f"[WARN] Image reading or OCR failed: {img_err}. Running binary string extraction...")
                text = extract_printable_strings(file_content)
                file_type = "image_extracted"

        # 6. Rich Text Format (RTF)
        elif ext == 'rtf':
            try:
                text = parse_rtf(file_content)
                file_type = "rtf"
            except Exception as rtf_err:
                print(f"[WARN] RTF parser failed: {rtf_err}. Running binary string extraction...")
                text = extract_printable_strings(file_content)
                file_type = "rtf_extracted"

        # 7. HTML / XML Web Markup
        elif ext in ['html', 'htm', 'xml']:
            try:
                text = parse_html(file_content)
                file_type = "web"
            except Exception as html_err:
                text = extract_printable_strings(file_content)
                file_type = "web_extracted"

        # 8. Excel Spreadsheets (XLSX / XLS)
        elif ext in ['xlsx', 'xls']:
            try:
                text = parse_xlsx(file_content)
                file_type = "excel"
            except Exception as xlsx_err:
                text = extract_printable_strings(file_content)
                file_type = "excel_extracted"

        # 9. PowerPoint Presentations (PPTX / PPT)
        elif ext in ['pptx', 'ppt']:
            try:
                text = parse_pptx(file_content)
                file_type = "presentation"
            except Exception as pptx_err:
                text = extract_printable_strings(file_content)
                file_type = "presentation_extracted"

        # 10. Default / Standard Text / Unknown Format Fallback
        else:
            print("[INFO] Unrecognized or standard text extension. Trying text decoders...")
            for encoding in ['utf-8', 'utf-16', 'latin-1']:
                try:
                    decoded = file_content.decode(encoding)
                    non_printable = sum(1 for c in decoded if ord(c) < 32 and c not in "\n\r\t")
                    if len(decoded) > 0 and (non_printable / len(decoded)) < 0.15:
                        text = decoded
                        file_type = f"text_{encoding}"
                        break
                except UnicodeDecodeError:
                    pass
            
            # Run printable strings extraction if decode yields gibberish or nothing
            if not text.strip():
                print("[INFO] Binary or unknown file format. Running smart string extraction...")
                text = extract_printable_strings(file_content)
                file_type = "binary_extracted"

        # Safe fallback: if after parsing we ended up with blank text, run string extraction
        if not text.strip():
            print("[WARN] Parsing yielded empty text. Running binary string extraction fallback...")
            text = extract_printable_strings(file_content)
            file_type = "fallback_extracted"

        return {
            "type": file_type,
            "content": text.strip()
        }

    except Exception as e:
        print(f"[ERROR] Parser Error: {e}")
        if "tesseract is not installed" in str(e).lower() or "find the file" in str(e).lower():
            print("[WARN] Tesseract command failed. Returning binary string extraction fallback...")
            text = extract_printable_strings(file_content)
            return {
                "type": "tesseract_fallback_extracted",
                "content": text.strip()
            }
        
        # Absolute bulletproof fallback
        try:
            fallback = extract_printable_strings(file_content)
            if fallback.strip():
                return {
                    "type": "error_fallback_extracted",
                    "content": fallback.strip()
                }
        except Exception:
            pass
            
        raise ValueError(f"Failed to read file: {str(e)}")